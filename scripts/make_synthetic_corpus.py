"""Generate synthetic corpus placeholders in tests/corpus/.

Resumes (deliberately parser-hostile, fictional data):

- synthetic_two_column.pdf  — two-column CSS layout (naive text extraction
  interleaves the columns), rendered with weasyprint
- synthetic_table_layout.docx — the whole resume laid out inside a table with
  merged cells and contact info in the page header, via python-docx

Job descriptions (realistic boilerplate-heavy text, fictional companies):

- synthetic_backend_jd.txt — senior backend role with explicit must-haves
- synthetic_ml_jd.txt — mid-level ML role with preferred qualifications

Run with: uv run python -m scripts.make_synthetic_corpus
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from weasyprint import HTML

CORPUS = Path(__file__).resolve().parent.parent / "tests" / "corpus"
CORPUS_RESUMES = CORPUS / "resumes"
CORPUS_JDS = CORPUS / "jds"

TWO_COLUMN_HTML = """
<html>
<head>
<style>
  @page { size: Letter; margin: 1.2cm; }
  body { font-family: Helvetica, sans-serif; font-size: 9.5pt; }
  .name { font-size: 20pt; letter-spacing: 4px; text-align: center; }
  .contact { text-align: center; font-size: 8pt; border-bottom: 2px solid #444; padding-bottom: 4px; }
  .columns { column-count: 2; column-gap: 24px; margin-top: 10px; }
  h2 { font-size: 10pt; text-transform: uppercase; border-bottom: 1px solid #999; }
  .job { margin-bottom: 6px; }
  .dates { float: right; font-style: italic; }
  ul { margin: 2px 0 6px 14px; padding: 0; }
</style>
</head>
<body>
  <div class="name">JORDAN&nbsp;RIVERA</div>
  <div class="contact">
    jordan.rivera@example.com | (555) 013-7788 | Austin, TX | linkedin.com/in/jordan-rivera-demo
  </div>
  <div class="columns">
    <h2>Experience</h2>
    <div class="job">
      <b>Data Engineer</b> <span class="dates">2021-08 – present</span><br>
      Cobalt Analytics, Austin TX
      <ul>
        <li>Maintain 40+ dbt models feeding exec dashboards</li>
        <li>Moved ingestion from cron to Dagster; 99.7% on-time runs</li>
        <li>Cut Snowflake spend 22% via clustering &amp; warehouse right-sizing</li>
      </ul>
    </div>
    <div class="job">
      <b>Analytics Engineer</b> <span class="dates">2019-02 – 2021-07</span><br>
      Brightpath Health
      <ul>
        <li>Built patient-flow reporting used by 3 hospital networks</li>
        <li>Automated HL7 feed validation in Python</li>
      </ul>
    </div>
    <h2>Education</h2>
    <p>BS Statistics, UT Austin, 2018-12</p>
    <h2>Skills</h2>
    <p>Python, SQL, dbt, Dagster, Snowflake, Airflow, Terraform, Looker</p>
    <h2>Projects</h2>
    <p><b>openlineage-viz</b> — graph viewer for OpenLineage events (Python, D3)</p>
    <h2>Certifications</h2>
    <p>SnowPro Core (2022-03)</p>
  </div>
</body>
</html>
"""


def make_two_column_pdf(out_dir: Path) -> Path:
    path = out_dir / "synthetic_two_column.pdf"
    HTML(string=TWO_COLUMN_HTML).write_pdf(str(path))
    return path


def make_table_layout_docx(out_dir: Path) -> Path:
    path = out_dir / "synthetic_table_layout.docx"
    doc = Document()

    # Contact info hidden in the page header — a classic extraction trap.
    header = doc.sections[0].header
    header.paragraphs[0].text = "SAM OKAFOR — sam.okafor@example.com — (555) 019-4452 — Chicago, IL"

    # Entire body laid out as a 2-column table with merged section cells.
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    title_cell = table.cell(0, 0).merge(table.cell(0, 1))
    run = title_cell.paragraphs[0].add_run("SENIOR BACKEND ENGINEER")
    run.bold = True
    run.font.size = Pt(16)

    table.cell(1, 0).text = "EXPERIENCE"
    exp = table.cell(1, 1)
    exp.text = "Nimbus Freight — Senior Backend Engineer (2020-05 – present)"
    for bullet in [
        "Own rate-quoting service: Go + Postgres, 8k rps peak",
        "Led move from EC2 to EKS across 11 services",
        "Reduced p99 quote latency from 900ms to 210ms",
    ]:
        exp.add_paragraph(f"• {bullet}")
    exp.add_paragraph("Parkside Labs — Backend Engineer (2017-09 – 2020-04)")
    exp.add_paragraph("• Built booking APIs in Django for 200k monthly users")

    table.cell(2, 0).text = "EDUCATION"
    table.cell(2, 1).text = "BEng Software Engineering, University of Illinois Chicago, 2017-05"

    table.cell(3, 0).text = "SKILLS"
    table.cell(3, 1).text = "Go • Python • Postgres • Kubernetes • gRPC • Kafka • Terraform"

    table.cell(4, 0).text = "AWARDS"
    table.cell(4, 1).text = "Nimbus 'Firefighter of the Year' 2022 (incident response)"

    doc.save(str(path))
    return path


BACKEND_JD = """\
Senior Backend Engineer — Corvid Logistics (Remote, US)

About us
Corvid Logistics builds routing and dispatch software for regional freight
carriers. Our platform handles 40M shipment events per day.

What you'll do
- Own services in our Python/FastAPI backend, from design through production
- Design PostgreSQL schemas and tune queries under heavy write load
- Build and operate event pipelines on Kafka
- Work with product to scope features; mentor two mid-level engineers

Minimum qualifications (required)
- 6+ years building production backend systems
- Expert-level Python; we use Python everywhere, and deep Python fluency is
  non-negotiable
- Strong PostgreSQL experience: schema design, indexing, query plans
- Experience running services on AWS (ECS or EKS)
- Clear written communication in a remote-first team

Nice to have
- Kafka or another event-streaming system
- Terraform, Datadog, or general observability experience
- Prior work in logistics, supply chain, or another operationally-heavy domain

We value engineers who take ownership, communicate proactively, and leave
codebases better than they found them.
"""

ML_JD = """\
Machine Learning Engineer — Halcyon Health (Hybrid, Boston)

Halcyon Health applies machine learning to clinical scheduling. We're a
40-person startup with real hospital customers and HIPAA obligations.

The role
You'll join a four-person ML team building no-show prediction and capacity
forecasting models, taking them from notebooks to production services.

You should have
- 2-4 years of experience training and shipping ML models in Python
- Solid grounding in scikit-learn and either PyTorch or TensorFlow
- SQL fluency and comfort working directly with messy healthcare data
- Experience deploying models behind APIs (we use FastAPI and Docker)

Preferred qualifications
- Experience with MLflow or another experiment-tracking tool
- Familiarity with HIPAA or other regulated-data environments
- Airflow or Dagster pipeline experience

We're a collaborative, low-ego team. You'll present model results to
clinicians, so the ability to explain technical tradeoffs to non-engineers
matters as much as modeling skill.
"""


def make_jds(out_dir: Path) -> list[Path]:
    paths = []
    for filename, body in [
        ("synthetic_backend_jd.txt", BACKEND_JD),
        ("synthetic_ml_jd.txt", ML_JD),
    ]:
        path = out_dir / filename
        path.write_text(body)
        paths.append(path)
    return paths


def main() -> None:
    CORPUS_RESUMES.mkdir(parents=True, exist_ok=True)
    CORPUS_JDS.mkdir(parents=True, exist_ok=True)
    outputs = [make_two_column_pdf(CORPUS_RESUMES), make_table_layout_docx(CORPUS_RESUMES)]
    outputs.extend(make_jds(CORPUS_JDS))
    for path in outputs:
        print(f"wrote {path}")


if __name__ == "__main__":
    main()
