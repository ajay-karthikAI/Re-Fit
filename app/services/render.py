"""Resume rendering: StructuredResume -> PDF/DOCX bytes."""

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Literal
import uuid

import anyio.to_thread
import pymupdf
from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from jinja2 import Environment, FileSystemLoader, select_autoescape
from sqlalchemy.ext.asyncio import AsyncSession
from weasyprint import HTML

from app.models import Document, DocumentKind, Profile, ResumeVersion
from app.schemas.render import RenderRequest, RenderResponse
from app.schemas.resume import StructuredResume
from app.services import storage
from app.services.errors import NotFoundError
from app.services.templates import (
    TemplateId,
    TemplateRenderSettings,
    normalize_template_variables,
    render_settings,
)

RenderFormat = Literal["pdf", "docx"]

DEFAULT_FONT_STACK = 'Aptos, Calibri, "Helvetica Neue", Arial, sans-serif'
DEFAULT_HEADING_FONT_STACK = 'Aptos Display, Aptos, Calibri, "Helvetica Neue", Arial, sans-serif'
DEFAULT_CONTACT_FONT_STACK = DEFAULT_FONT_STACK
DEFAULT_ACCENT_COLOR = "#24526b"
DEFAULT_MARGIN_IN = 0.7
PDF_FONT_ATTEMPTS = (11.0, 10.5, 10.0)

TEMPLATE_ROOT = Path(__file__).resolve().parents[2] / "templates" / "resume"


@dataclass(frozen=True)
class RenderedArtifact:
    content: bytes
    page_count: int | None
    warnings: list[str]
    template: TemplateId
    format: RenderFormat
    font_size: float | None = None


class TemplateResumeRenderer:
    def __init__(
        self,
        template: TemplateId = "classic",
        template_variables: dict[str, object] | None = None,
    ) -> None:
        self.template = template
        self.template_variables = template_variables

    def render_pdf(self, resume: StructuredResume) -> bytes:
        return render_pdf(
            resume, template=self.template, template_variables=self.template_variables
        )


class ClassicResumeRenderer:
    def render_pdf(self, resume: StructuredResume) -> bytes:
        return render_pdf(resume)


def _environment() -> Environment:
    return Environment(
        loader=FileSystemLoader(TEMPLATE_ROOT),
        autoescape=select_autoescape(("html", "j2")),
        trim_blocks=True,
        lstrip_blocks=True,
    )


def _contact_line(resume: StructuredResume) -> str:
    values = [
        resume.contact.email,
        resume.contact.phone,
        resume.contact.location,
        str(resume.contact.linkedin_url or ""),
        str(resume.contact.github_url or ""),
        str(resume.contact.portfolio_url or ""),
    ]
    return " | ".join(value for value in values if value)


def _format_date_range(start_date: str, end_date: str | None) -> str:
    return f"{start_date} - {end_date or 'Present'}"


def _markdown_paragraphs(markdown: str) -> list[str]:
    return [paragraph.strip() for paragraph in markdown.split("\n\n") if paragraph.strip()]


def _render_html(
    resume: StructuredResume,
    *,
    template: TemplateId,
    base_font_size: float,
    show_page_numbers: bool,
    font_stack: str,
    heading_font_stack: str,
    contact_font_stack: str,
    accent_color: str,
    margin_in: float,
) -> str:
    jinja_template = _environment().get_template(f"{template}/resume.html.j2")
    return jinja_template.render(
        resume=resume,
        contact_line=_contact_line(resume),
        format_date_range=_format_date_range,
        base_font_size=base_font_size,
        show_page_numbers=show_page_numbers,
        font_stack=font_stack,
        heading_font_stack=heading_font_stack,
        contact_font_stack=contact_font_stack,
        accent_color=accent_color,
        margin_in=margin_in,
    )


def _render_cover_letter_html(
    resume: StructuredResume,
    body_markdown: str,
    *,
    template: TemplateId,
    base_font_size: float,
    show_page_numbers: bool,
    font_stack: str,
    heading_font_stack: str,
    contact_font_stack: str,
    accent_color: str,
    margin_in: float,
) -> str:
    jinja_template = _environment().get_template(f"{template}/cover_letter.html.j2")
    return jinja_template.render(
        resume=resume,
        contact_line=_contact_line(resume),
        paragraphs=_markdown_paragraphs(body_markdown),
        base_font_size=base_font_size,
        show_page_numbers=show_page_numbers,
        font_stack=font_stack,
        heading_font_stack=heading_font_stack,
        contact_font_stack=contact_font_stack,
        accent_color=accent_color,
        margin_in=margin_in,
    )


def _html_to_pdf(html: str) -> bytes:
    return HTML(string=html, base_url=str(TEMPLATE_ROOT)).write_pdf()


def _pdf_page_count(pdf_bytes: bytes) -> int:
    with pymupdf.open(stream=pdf_bytes, filetype="pdf") as doc:
        return doc.page_count


def _template_variables(
    template: TemplateId,
    template_variables: dict[str, object] | None,
    *,
    accent_color: str | None = None,
    font_scale: float | None = None,
) -> dict[str, str | float]:
    variables = dict(template_variables or {})
    if accent_color is not None:
        variables["accent_color"] = accent_color
    if font_scale is not None:
        variables["font_scale"] = font_scale
    return normalize_template_variables(template, variables)


def _settings(
    template: TemplateId,
    template_variables: dict[str, object] | None,
    *,
    accent_color: str | None = None,
    font_scale: float | None = None,
) -> TemplateRenderSettings:
    return render_settings(
        template,
        _template_variables(
            template,
            template_variables,
            accent_color=accent_color,
            font_scale=font_scale,
        ),
    )


def render_pdf_result(
    resume: StructuredResume,
    template: TemplateId = "classic",
    *,
    template_variables: dict[str, object] | None = None,
    font_stack: str | None = None,
    heading_font_stack: str | None = None,
    accent_color: str | None = None,
    font_scale: float | None = None,
    margin_in: float | None = None,
) -> RenderedArtifact:
    settings = _settings(
        template,
        template_variables,
        accent_color=accent_color,
        font_scale=font_scale,
    )
    chosen_font = font_stack or settings.font_stack
    chosen_heading = heading_font_stack or settings.heading_font_stack
    chosen_contact = settings.contact_font_stack
    chosen_margin = margin_in if margin_in is not None else settings.margin_in
    warnings: list[str] = []
    final_pdf = b""
    final_page_count = 0
    final_font_size = settings.pdf_font_attempts[-1]

    for font_size in settings.pdf_font_attempts:
        html = _render_html(
            resume,
            template=template,
            base_font_size=font_size,
            show_page_numbers=False,
            font_stack=chosen_font,
            heading_font_stack=chosen_heading,
            contact_font_stack=chosen_contact,
            accent_color=settings.accent_color,
            margin_in=chosen_margin,
        )
        pdf = _html_to_pdf(html)
        page_count = _pdf_page_count(pdf)
        show_page_numbers = page_count > 1
        if show_page_numbers:
            html = _render_html(
                resume,
                template=template,
                base_font_size=font_size,
                show_page_numbers=True,
                font_stack=chosen_font,
                heading_font_stack=chosen_heading,
                contact_font_stack=chosen_contact,
                accent_color=settings.accent_color,
                margin_in=chosen_margin,
            )
            pdf = _html_to_pdf(html)
            page_count = _pdf_page_count(pdf)

        final_pdf = pdf
        final_page_count = page_count
        final_font_size = font_size
        if page_count <= 2:
            if font_size != settings.pdf_font_attempts[0]:
                warnings.append(
                    f"autofit reduced base font size from {settings.pdf_font_attempts[0]:g}pt "
                    f"to {font_size:g}pt"
                )
            break

    if final_page_count > 2:
        warnings.append(
            "render exceeds 2 pages after autofit; review content density or template settings"
        )

    return RenderedArtifact(
        content=final_pdf,
        page_count=final_page_count,
        warnings=warnings,
        template=template,
        format="pdf",
        font_size=final_font_size,
    )


def render_pdf(
    resume: StructuredResume,
    template: TemplateId = "classic",
    *,
    template_variables: dict[str, object] | None = None,
) -> bytes:
    return render_pdf_result(
        resume,
        template=template,
        template_variables=template_variables,
    ).content


def render_cover_letter_pdf_result(
    resume: StructuredResume,
    body_markdown: str,
    template: TemplateId = "classic",
    *,
    template_variables: dict[str, object] | None = None,
    font_stack: str | None = None,
    heading_font_stack: str | None = None,
    accent_color: str | None = None,
    font_scale: float | None = None,
    margin_in: float | None = None,
) -> RenderedArtifact:
    settings = _settings(
        template,
        template_variables,
        accent_color=accent_color,
        font_scale=font_scale,
    )
    chosen_font = font_stack or settings.font_stack
    chosen_heading = heading_font_stack or settings.heading_font_stack
    chosen_contact = settings.contact_font_stack
    chosen_margin = margin_in if margin_in is not None else settings.margin_in
    warnings: list[str] = []

    html = _render_cover_letter_html(
        resume,
        body_markdown,
        template=template,
        base_font_size=11.0,
        show_page_numbers=False,
        font_stack=chosen_font,
        heading_font_stack=chosen_heading,
        contact_font_stack=chosen_contact,
        accent_color=settings.accent_color,
        margin_in=chosen_margin,
    )
    pdf = _html_to_pdf(html)
    page_count = _pdf_page_count(pdf)
    if page_count > 1:
        html = _render_cover_letter_html(
            resume,
            body_markdown,
            template=template,
            base_font_size=11.0,
            show_page_numbers=True,
            font_stack=chosen_font,
            heading_font_stack=chosen_heading,
            contact_font_stack=chosen_contact,
            accent_color=settings.accent_color,
            margin_in=chosen_margin,
        )
        pdf = _html_to_pdf(html)
        page_count = _pdf_page_count(pdf)
    if page_count > 2:
        warnings.append("cover letter exceeds 2 pages; review prose length")
    return RenderedArtifact(
        content=pdf,
        page_count=page_count,
        warnings=warnings,
        template=template,
        format="pdf",
        font_size=settings.pdf_font_attempts[0],
    )


def _set_style_font(
    doc: DocxDocument,
    style_name: str,
    *,
    font_name: str,
    font_size_pt: float,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    style = doc.styles[style_name]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size_pt)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = color


def _prepare_docx_styles(doc: DocxDocument, settings: TemplateRenderSettings) -> None:
    accent = RGBColor.from_string(settings.accent_color.lstrip("#").upper())
    styles = doc.styles
    if "Resume Contact" not in styles:
        styles.add_style("Resume Contact", WD_STYLE_TYPE.PARAGRAPH)
    if "Resume Metadata" not in styles:
        styles.add_style("Resume Metadata", WD_STYLE_TYPE.PARAGRAPH)

    _set_style_font(
        doc, "Normal", font_name=settings.docx_body_font, font_size_pt=settings.docx_body_size
    )
    _set_style_font(
        doc,
        "Heading 1",
        font_name=settings.docx_heading_font,
        font_size_pt=settings.docx_heading_size,
        bold=True,
        color=accent,
    )
    _set_style_font(
        doc,
        "Heading 2",
        font_name=settings.docx_body_font,
        font_size_pt=settings.docx_subheading_size,
        bold=True,
    )
    _set_style_font(
        doc,
        "Resume Contact",
        font_name=settings.docx_contact_font,
        font_size_pt=settings.docx_contact_size,
    )
    _set_style_font(
        doc,
        "Resume Metadata",
        font_name=settings.docx_body_font,
        font_size_pt=settings.docx_contact_size,
    )

    for style_name in ["Normal", "Heading 1", "Heading 2", "Resume Contact", "Resume Metadata"]:
        paragraph_format = styles[style_name].paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(settings.docx_heading_after)
        paragraph_format.line_spacing = settings.docx_line_spacing

    styles["Heading 1"].paragraph_format.space_before = Pt(settings.docx_section_before)
    styles["Heading 1"].paragraph_format.space_after = Pt(settings.docx_heading_after)
    styles["Heading 2"].paragraph_format.space_before = Pt(settings.docx_heading_after)
    styles["Heading 2"].paragraph_format.space_after = Pt(1)


def _add_bullet(doc: DocxDocument, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(2)


def render_docx_result(
    resume: StructuredResume,
    template: TemplateId = "classic",
    *,
    template_variables: dict[str, object] | None = None,
    accent_color: str | None = None,
    font_scale: float | None = None,
) -> RenderedArtifact:
    settings = _settings(
        template,
        template_variables,
        accent_color=accent_color,
        font_scale=font_scale,
    )

    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Inches(settings.margin_in)
    section.bottom_margin = Inches(settings.margin_in)
    section.left_margin = Inches(settings.margin_in)
    section.right_margin = Inches(settings.margin_in)
    _prepare_docx_styles(doc, settings)

    name = doc.add_heading(resume.contact.full_name, level=1)
    name.style = doc.styles["Heading 1"]
    contact = doc.add_paragraph(_contact_line(resume), style="Resume Contact")
    contact.paragraph_format.space_after = Pt(8)

    if resume.summary:
        doc.add_heading("SUMMARY", level=1)
        doc.add_paragraph(resume.summary)

    if resume.experience:
        doc.add_heading("EXPERIENCE", level=1)
        for item in resume.experience:
            heading = doc.add_heading(level=2)
            heading.add_run(f"{item.title} - {item.company}")
            dates = doc.add_paragraph(
                _format_date_range(item.start_date, item.end_date), style="Resume Metadata"
            )
            if item.location:
                dates.add_run(f" | {item.location}")
            for bullet in item.bullets:
                _add_bullet(doc, bullet)
            if item.technologies:
                doc.add_paragraph(f"Technologies: {', '.join(item.technologies)}")

    if resume.education:
        doc.add_heading("EDUCATION", level=1)
        for item in resume.education:
            heading = doc.add_heading(level=2)
            heading.add_run(
                f"{item.degree}{', ' + item.field if item.field else ''} - {item.institution}"
            )
            if item.graduation_date:
                doc.add_paragraph(item.graduation_date, style="Resume Metadata")
            for detail in item.details:
                _add_bullet(doc, detail)

    if resume.skills:
        doc.add_heading("SKILLS", level=1)
        for group in resume.skills:
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{group.category}: ").bold = True
            paragraph.add_run(", ".join(group.items))

    if resume.projects:
        doc.add_heading("PROJECTS", level=1)
        for project in resume.projects:
            heading = doc.add_heading(project.name, level=2)
            if project.url:
                heading.add_run(f" - {project.url}")
            if project.description:
                doc.add_paragraph(project.description)
            for bullet in project.bullets:
                _add_bullet(doc, bullet)
            if project.technologies:
                doc.add_paragraph(f"Technologies: {', '.join(project.technologies)}")

    buffer = BytesIO()
    doc.save(buffer)
    return RenderedArtifact(
        content=buffer.getvalue(),
        page_count=None,
        warnings=[],
        template=template,
        format="docx",
        font_size=None,
    )


def render_docx(
    resume: StructuredResume,
    template: TemplateId = "classic",
    *,
    template_variables: dict[str, object] | None = None,
) -> bytes:
    return render_docx_result(
        resume,
        template=template,
        template_variables=template_variables,
    ).content


def render_cover_letter_docx_result(
    resume: StructuredResume,
    body_markdown: str,
    template: TemplateId = "classic",
    *,
    template_variables: dict[str, object] | None = None,
    accent_color: str | None = None,
    font_scale: float | None = None,
) -> RenderedArtifact:
    settings = _settings(
        template,
        template_variables,
        accent_color=accent_color,
        font_scale=font_scale,
    )

    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Inches(settings.margin_in)
    section.bottom_margin = Inches(settings.margin_in)
    section.left_margin = Inches(settings.margin_in)
    section.right_margin = Inches(settings.margin_in)
    _prepare_docx_styles(doc, settings)

    name = doc.add_heading(resume.contact.full_name, level=1)
    name.style = doc.styles["Heading 1"]
    contact = doc.add_paragraph(_contact_line(resume), style="Resume Contact")
    contact.paragraph_format.space_after = Pt(16)

    for paragraph_text in _markdown_paragraphs(body_markdown):
        paragraph = doc.add_paragraph(paragraph_text)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.12

    buffer = BytesIO()
    doc.save(buffer)
    return RenderedArtifact(
        content=buffer.getvalue(),
        page_count=None,
        warnings=[],
        template=template,
        format="docx",
        font_size=None,
    )


def _payload_template_choice(
    version: ResumeVersion,
    payload: RenderRequest,
) -> tuple[TemplateId, dict[str, str | float], bool]:
    template_id = payload.template or version.template_id or "classic"
    variables = dict(version.template_variables or {})
    should_persist = False
    if payload.template is not None:
        variables = {}
        should_persist = True
    if payload.template_variables is not None:
        variables.update(payload.template_variables)
        should_persist = True
    if payload.accent_color is not None:
        variables["accent_color"] = payload.accent_color
        should_persist = True
    if payload.font_scale is not None:
        variables["font_scale"] = payload.font_scale
        should_persist = True
    normalized = normalize_template_variables(template_id, variables)
    return template_id, normalized, should_persist


async def render_version_document(
    session: AsyncSession,
    version_id: uuid.UUID,
    payload: RenderRequest,
    *,
    expires_in: int = 900,
) -> RenderResponse:
    version = await session.get(ResumeVersion, version_id)
    if version is None or version.deleted_at is not None:
        raise NotFoundError(f"resume version {version_id} not found")
    profile = await session.get(Profile, version.profile_id)
    if profile is None:
        raise NotFoundError(f"profile {version.profile_id} not found")

    resume = StructuredResume.model_validate(version.data)
    template_id, template_variables, should_persist_template = _payload_template_choice(
        version,
        payload,
    )
    if should_persist_template:
        version.template_id = template_id
        version.template_variables = template_variables
        await session.flush()

    if payload.format == "pdf":
        artifact = await anyio.to_thread.run_sync(
            lambda: render_pdf_result(
                resume,
                template=template_id,
                template_variables=template_variables,
                font_stack=payload.font,
                heading_font_stack=payload.font,
            )
        )
        kind = DocumentKind.resume_pdf
        ext = "pdf"
        content_type = "application/pdf"
    else:
        artifact = await anyio.to_thread.run_sync(
            lambda: render_docx_result(
                resume,
                template=template_id,
                template_variables=template_variables,
            )
        )
        kind = DocumentKind.resume_docx
        ext = "docx"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    s3_key = f"documents/{profile.user_id}/{version.id}/{kind.value}.{ext}"
    await anyio.to_thread.run_sync(storage.put_object, s3_key, artifact.content, content_type)

    document = Document(resume_version_id=version.id, kind=kind, s3_key=s3_key)
    session.add(document)
    await session.commit()
    await session.refresh(document)

    download_url = await anyio.to_thread.run_sync(storage.presigned_get_url, s3_key, expires_in)
    return RenderResponse(
        document_id=document.id,
        kind=kind.value,
        s3_key=s3_key,
        download_url=download_url,
        expires_in=expires_in,
        page_count=artifact.page_count,
        warnings=artifact.warnings,
    )
